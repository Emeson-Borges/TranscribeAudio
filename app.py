import streamlit as st
import io
import soundfile as sf
import speech_recognition as sr
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pyperclip
import webbrowser
import pygetwindow as gw
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from docx.shared import Pt, RGBColor


# Função para transcrever áudio para texto
def transcrever_audio(audio_data, audio_format):
    r = sr.Recognizer()
    with io.BytesIO(audio_data) as audio_io:
        audio_io.seek(0)
        with sf.SoundFile(audio_io) as audio_file:
            audio_data = audio_file.read(dtype='float32')
            audio_rate = audio_file.samplerate
            audio_channels = audio_file.channels
            # Converter para PCM WAV se não estiver no formato adequado
            if audio_format != 'wav':
                with io.BytesIO() as wav_io:
                    sf.write(wav_io, audio_data, audio_rate, format='WAV', subtype='PCM_24')
                    wav_io.seek(0)
                    audio_data = wav_io.read()
        with io.BytesIO(audio_data) as source:
            with sr.AudioFile(source) as audio:
                audio_data = r.record(audio)
                text = r.recognize_google(audio_data, language='pt-BR')  # Altere o idioma conforme necessário
    return text


# Função para encontrar a janela do WhatsApp
def get_whatsapp_window():
    whatsapp_windows = gw.getWindowsWithTitle('WhatsApp')
    for window in whatsapp_windows:
        if "web.whatsapp.com" in window.url:
            return window
    return None


# Função para enviar texto para o WhatsApp
def enviar_para_whatsapp(texto):
    whatsapp_window = get_whatsapp_window()
    if whatsapp_window:
        whatsapp_window.activate()
    else:
        webbrowser.open("https://web.whatsapp.com/")
    pyperclip.copy(texto)


# Funções para salvar o texto em diferentes formatos
def save_as_word(texto, filename):
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run(texto)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = False
    font.italic = False
    font.underline = False
    font.color.rgb = RGBColor(0, 0, 0)
    document.save(filename + ".docx")
    st.success(f"Texto salvo em '{filename}.docx'")


def save_as_pdf(texto, filename):
    filename += ".pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    style_normal = styles["Normal"]
    text_lines = texto.split("\n")
    formatted_text = [Paragraph(line, style_normal) for line in text_lines]
    doc.build(formatted_text)
    st.success(f"Texto salvo em '{filename}'")


def save_as_txt(texto, filename):
    with open(filename + ".txt", "w", encoding="utf-8") as f:
        f.write(texto)
    st.success(f"Texto salvo em '{filename}.txt'")


# Função principal
def main():
    st.title("Transcrição de Áudio para Texto")

    # Sidebar para upload do arquivo de áudio
    st.sidebar.title("Upload do Áudio")
    uploaded_file = st.sidebar.file_uploader("Selecione um arquivo de áudio")

    # # Campo de entrada para nome do arquivo
    # filename = st.sidebar.text_input("Nome do arquivo")

    # Inicializa a variável de sessão para armazenar o texto transcrito
    if 'texto_transcrito' not in st.session_state:
        st.session_state.texto_transcrito = None

    # Processamento do arquivo de áudio
    if uploaded_file is not None:
        # Lê o arquivo de áudio
        audio_data = uploaded_file.getvalue()
        audio_format = uploaded_file.name.split(".")[-1]

        # Exibe o arquivo de áudio carregado
        st.audio(io.BytesIO(audio_data), format='audio/' + audio_format)

        # Botão para transcrever o áudio
        if st.sidebar.button("Transcrever"):
            st.session_state.texto_transcrito = None
            with st.spinner("Transcrevendo áudio..."):
                try:
                    # Transcreve o áudio
                    texto_transcrito = transcrever_audio(audio_data, audio_format)
                    # Armazena o texto transcrito na variável de sessão
                    st.session_state.texto_transcrito = texto_transcrito
                    st.sidebar.success("Transcrição concluída com sucesso!")
                except Exception as e:
                    st.sidebar.error(f"Erro durante a transcrição: {str(e)}")

    # Exibe a mensagem de andamento ou de conclusão da transcrição
    # if st.session_state.texto_transcrito is None:
    #     st.sidebar.info("Transcrição em andamento...")
    # else:
    #     st.sidebar.success("Transcrição concluída com sucesso!")

        # Exibe o texto transcrito
        st.header("Texto Transcrito:")
        st.write(st.session_state.texto_transcrito)

        # Se houver texto transcrito, exibe os botões de ação
        if st.session_state.texto_transcrito is not None:
             # Campo de entrada para nome do arquivo
            filename = st.sidebar.text_input("Nome do arquivo")
            # Botão para enviar para o WhatsApp
            if st.sidebar.button("Enviar para WhatsApp"):
                enviar_para_whatsapp(st.session_state.texto_transcrito)

            # Seletor de formato para salvar o texto
            formato = st.sidebar.selectbox("Selecione o formato para salvar o texto:", ("Word", "PDF", "TXT"))

            # Botões para download do arquivo
            if formato == "Word":
                if st.sidebar.button("Baixar Word"):
                    save_as_word(st.session_state.texto_transcrito, filename)
            elif formato == "PDF":
                if st.sidebar.button("Baixar PDF"):
                    save_as_pdf(st.session_state.texto_transcrito, filename)
            elif formato == "TXT":
                if st.sidebar.button("Baixar TXT"):
                    save_as_txt(st.session_state.texto_transcrito, filename)


if __name__ == "__main__":
    main()
